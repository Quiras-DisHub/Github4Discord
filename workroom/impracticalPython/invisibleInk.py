import docx
from docx.shared import RGBColor, Pt

fakeText = docx.Document("fakeMMessage.docx")
fakeList = []
for paragraph in fakeText.paragraphs:
    fakeList.appent(paragraph.text)

realText = docx.Document("realMessage.docx")
realList = []
for paragraph in realText.paragraphs:
    if len(paragraph.text) != 0:
        realList.append(paragraph.text)

doc = docx.Document("template.docx")

doc.add_heading("Morland Holmes", 0)
subtitle = doc.add_heading("Global Consulting & Negotioations", 1)
subtitle.alignment = 1
doc.add_heading("", 1)
doc.add_paragraph("December 17, 2015")
doc.add_paragraph("")

